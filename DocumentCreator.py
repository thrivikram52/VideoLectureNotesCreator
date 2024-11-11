import os
from pathlib import Path
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

class DocumentCreator:
    def __init__(self):
        # Configure logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def create_document(self, output_folder):
        """
        Converts lecture_notes.md to Word format and appends lecture slides
        
        Args:
            output_folder (str): Path to the folder containing lecture_notes.md
            
        Returns:
            str: Path to the generated Word document
        """
        try:
            md_file_path = os.path.join(output_folder, "lecture_notes.md")
            docx_file_path = os.path.join(output_folder, "lecture_notes.docx")
            
            # Convert markdown to docx using pandoc with no-bookmarks option
            subprocess.run([
                'pandoc',
                md_file_path,
                '-o', docx_file_path,
                '--from=markdown-auto_identifiers',  # This removes the automatic bookmarks
                '--to=docx'
            ], check=True)
            
            # Now add the slides to the existing document
            doc = Document(docx_file_path)
            
            # Add page break before slides
            doc.add_page_break()
            doc.add_heading('Lecture Slides', 1)
            
            # Add slides (existing code)
            image_files = sorted([f for f in os.listdir(output_folder) if f.endswith('.png')])
            if image_files:
                for img_file in image_files:
                    img_path = os.path.join(output_folder, img_file)
                    doc.add_picture(img_path, width=Inches(6.0))
                    caption = doc.add_paragraph(f"Slide: {img_file}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

            doc.save(docx_file_path)
            return docx_file_path

        except subprocess.CalledProcessError as e:
            self.logger.error(f"Pandoc conversion failed: {str(e)}")
            raise
        except Exception as e:
            self.logger.error(f"Error creating Word document: {str(e)}")
            raise

    def apply_document_styling(self, doc):
        """
        Apply consistent styling to the Word document
        
        Args:
            doc: Document object to style
        """
        # Style the default paragraph font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Style the headings
        for i in range(1, 4):  # Heading levels 1-3
            style = doc.styles[f'Heading {i}']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(14 - i)  # Decreasing size for each level
            font.bold = True

    def add_table_of_contents(self, doc):
        """Add a table of contents to the document"""
        doc.add_paragraph("Table of Contents")
        doc.add_paragraph("").add_run().add_break()
        # Note: Word will need to update the TOC when opened

    def add_header_footer(self, doc, header_text="Lecture Notes"):
        """Add header and footer to the document"""
        section = doc.sections[0]
        header = section.header
        footer = section.footer
        
        # Add header
        header_para = header.paragraphs[0]
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add footer with page numbers
        footer_para = footer.paragraphs[0]
        footer_para.text = "Page "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run().add_field('PAGE')
